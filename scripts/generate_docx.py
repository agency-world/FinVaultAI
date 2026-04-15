#!/usr/bin/env python3
"""
FinVault AI — Technical Guide Generator
Creates a comprehensive DOCX technical guide with embedded screenshots.
"""

import os
import sys
from pathlib import Path

# Ensure we can import docx
sys.path.insert(0, "/opt/homebrew/lib/python3.14/site-packages")

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT

DOCS_DIR = Path(__file__).resolve().parent.parent / "docs"
SCREENSHOTS = DOCS_DIR / "screenshots"
OUTPUT = DOCS_DIR / "FinVault_AI_Technical_Guide.docx"


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), color_hex)
    shading_elm.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_heading_styled(doc, text, level=1):
    """Add a heading with indigo color."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        if level == 1:
            run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
            run.font.size = Pt(24)
        elif level == 2:
            run.font.color.rgb = RGBColor(0x63, 0x66, 0xF1)
            run.font.size = Pt(18)
        elif level == 3:
            run.font.color.rgb = RGBColor(0x63, 0x66, 0xF1)
            run.font.size = Pt(14)
    return h


def add_para(doc, text, bold=False, italic=False, size=11, color=None, align=None, space_after=6):
    """Add a styled paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if color:
        run.font.color.rgb = color
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    from docx.shared import Pt as PtShared
    p.paragraph_format.space_after = PtShared(space_after)
    return p


def add_bullet(doc, text, level=0):
    """Add a bullet point."""
    p = doc.add_paragraph(text, style="List Bullet")
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.5 * level)
    return p


def add_screenshot(doc, filename, caption="", width=5.8):
    """Add a screenshot image with caption."""
    img_path = SCREENSHOTS / filename
    if img_path.exists():
        doc.add_picture(str(img_path), width=Inches(width))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cap.add_run(caption)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
            run.italic = True
    else:
        add_para(doc, f"[Screenshot: {filename}]", italic=True,
                 color=RGBColor(0x94, 0xA3, 0xB8), align="center")


def add_info_table(doc, rows_data, col_widths=None):
    """Add a styled two-column info table."""
    table = doc.add_table(rows=len(rows_data), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(rows_data):
        cell_l = table.cell(i, 0)
        cell_r = table.cell(i, 1)
        cell_l.text = label
        cell_r.text = value
        # Style
        for p in cell_l.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
                r.font.name = "Calibri"
        for p in cell_r.paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
                r.font.name = "Calibri"
        if i % 2 == 0:
            set_cell_shading(cell_l, "F0F0FF")
            set_cell_shading(cell_r, "F0F0FF")
    return table


def add_code_block(doc, code_text):
    """Add a monospaced code block."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x1B, 0x4B)
    return p


def build_document():
    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ── Default font ──
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    # ════════════════════════════════════════════════════════════
    # COVER PAGE
    # ════════════════════════════════════════════════════════════
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("FinVault AI")
    run.font.size = Pt(42)
    run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
    run.bold = True
    run.font.name = "Calibri"

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Technical Guide & Product Documentation")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x63, 0x66, 0xF1)
    run.font.name = "Calibri"

    add_para(doc, "", size=6)

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run("On-Device Financial Intelligence. Private by Design.")
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    run.italic = True

    for _ in range(6):
        doc.add_paragraph()

    meta_items = [
        ("Version", "3.0"),
        ("Date", "April 2026"),
        ("Classification", "Confidential"),
        ("Engine", "Google Gemma 4 via LM Studio"),
    ]
    for label, value in meta_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f"{label}: ")
        r1.bold = True
        r1.font.size = Pt(10)
        r1.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        r2 = p.add_run(value)
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ════════════════════════════════════════════════════════════
    add_heading_styled(doc, "Table of Contents", level=1)
    toc_items = [
        "1. Executive Summary",
        "2. Product Overview",
        "3. System Architecture",
        "4. Development Guide",
        "5. Agent Engineering & LLM Configuration",
        "6. Security Architecture",
        "7. Deployment Guide",
        "8. User Demo Guide",
        "9. Product Roadmap",
        "10. Appendix",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # 1. EXECUTIVE SUMMARY
    # ════════════════════════════════════════════════════════════
    add_heading_styled(doc, "1. Executive Summary", level=1)

    add_para(doc, (
        "FinVault AI is a secure, on-device financial intelligence platform that delivers "
        "AI-powered accounting analysis without transmitting a single byte of data to external servers. "
        "Built on Google's Gemma 4 large language model running locally via LM Studio, FinVault AI "
        "provides three core capabilities: intelligent report generation, natural language data querying, "
        "and automated document OCR processing."
    ))
    add_para(doc, (
        "The platform is designed for financial professionals who handle sensitive accounting data "
        "and require enterprise-grade AI capabilities while maintaining complete data sovereignty. "
        "FinVault AI operates in air-gapped environments, requires no internet connection, "
        "and processes all data exclusively within the user's local machine."
    ))

    add_heading_styled(doc, "Key Highlights", level=2)
    highlights = [
        "100% local processing — zero cloud dependency",
        "Google Gemma 4 (google/gemma-4-e4b) for state-of-the-art inference",
        "SOX, GDPR, and CCPA compliant by architecture",
        "Three integrated AI modules: Reports, Query, OCR",
        "Sub-second inference for most financial queries",
        "Air-gap compatible — works without internet",
    ]
    for h in highlights:
        add_bullet(doc, h)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # 2. PRODUCT OVERVIEW
    # ════════════════════════════════════════════════════════════
    add_heading_styled(doc, "2. Product Overview", level=1)

    add_heading_styled(doc, "2.1 Report Generator", level=2)
    add_para(doc, (
        "The Report Generator transforms raw accounting data (CSV/Excel) into professional "
        "financial reports. Users select from seven report types including Executive Summary, "
        "Expense Breakdown, Revenue Analysis, Budget Variance, Cash Flow, Accounts Payable/Receivable, "
        "and Custom Analysis. The AI analyzes the data, computes key metrics, and produces "
        "narratives with actionable insights."
    ))
    add_screenshot(doc, "01_dashboard_report.png", "Figure 1: Report Generator with data preview, metrics, and chart visualization")

    add_heading_styled(doc, "2.2 Natural Language Data Query", level=2)
    add_para(doc, (
        "The Data Query module provides a conversational interface for financial data analysis. "
        "Users ask questions in plain English (e.g., 'What is Q1 revenue?', 'Which category has "
        "the highest expenses?') and receive precise answers with step-by-step calculations. "
        "The chat maintains conversation context for follow-up questions and displays column "
        "metadata for transparency."
    ))
    add_screenshot(doc, "02_data_query.png", "Figure 2: Natural Language Query interface with chat history and column chips")

    add_heading_styled(doc, "2.3 Document OCR", level=2)
    add_para(doc, (
        "The OCR module processes scanned financial documents — invoices, work orders, receipts, "
        "and purchase orders. Tesseract 5.5 performs local text extraction with confidence scoring, "
        "then Gemma 4 parses structured fields (invoice numbers, dates, amounts, line items). "
        "A vision-based fallback is available when Tesseract is not installed."
    ))
    add_screenshot(doc, "03_ocr_processing.png", "Figure 3: OCR processing with confidence metrics and parsed structured data")

    add_heading_styled(doc, "2.4 Admin Panel", level=2)
    add_para(doc, (
        "The sidebar Admin section displays real-time model information (Gemma 4 model card), "
        "inference configuration (temperature, max tokens, top-p), and system health status "
        "(OCR engine availability, LLM server connectivity)."
    ))
    add_screenshot(doc, "04_admin_panel.png", "Figure 4: Admin panel showing model info, inference settings, and system health")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # 3. SYSTEM ARCHITECTURE
    # ════════════════════════════════════════════════════════════
    add_heading_styled(doc, "3. System Architecture", level=1)

    add_heading_styled(doc, "3.1 Architecture Overview", level=2)
    add_para(doc, (
        "FinVault AI follows a strictly local architecture where all components run on the user's "
        "machine. The system consists of four layers: the Streamlit web UI, the Python application "
        "layer, the LM Studio inference server, and the local Tesseract OCR engine."
    ))
    add_screenshot(doc, "05_architecture.png", "Figure 5: System architecture — all processing stays local")

    add_heading_styled(doc, "3.2 Component Stack", level=2)
    components = [
        ("Frontend", "Streamlit 1.x with custom dark theme (indigo accent #6366F1)"),
        ("Application", "Python 3.14 — modular structure with core/, ui/, utils/, config/"),
        ("LLM Engine", "Google Gemma 4 (google/gemma-4-e4b) via LM Studio localhost:1234"),
        ("OCR Engine", "Tesseract 5.5 (Homebrew) with PIL preprocessing pipeline"),
        ("API Protocol", "OpenAI-compatible REST API (openai Python SDK → localhost)"),
        ("Data Format", "CSV, Excel (xlsx/xls), PNG, JPG, PDF"),
        ("Charts", "Plotly Express for interactive financial visualizations"),
    ]
    add_info_table(doc, components)

    add_heading_styled(doc, "3.3 Data Flow", level=2)
    flow_steps = [
        "User uploads accounting data (CSV/Excel) or scanned document (image/PDF) via browser",
        "Streamlit receives file into session memory (never written to disk)",
        "Application layer preprocesses data: pandas for tabular, PIL for images",
        "For OCR: Tesseract extracts text locally, returns confidence + word count",
        "Prompt is constructed with system context + user data + instructions",
        "Request sent to LM Studio at http://localhost:1234/v1/chat/completions",
        "Gemma 4 processes request entirely on-device (GPU/CPU inference)",
        "Response streamed back to application layer",
        "Results displayed in Streamlit UI with download options",
        "Session ends → all data purged from memory. Nothing persists.",
    ]
    for i, step in enumerate(flow_steps, 1):
        add_bullet(doc, f"Step {i}: {step}")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # 4. DEVELOPMENT GUIDE
    # ════════════════════════════════════════════════════════════
    add_heading_styled(doc, "4. Development Guide", level=1)

    add_heading_styled(doc, "4.1 Project Structure", level=2)
    add_code_block(doc, """FinVault AI/
├── app.py                    # Main Streamlit entry point
├── config/
│   └── settings.py           # All constants, prompts, defaults
├── core/
│   ├── llm_client.py         # OpenAI SDK wrapper for LM Studio
│   ├── report_generator.py   # Report prompt construction
│   ├── data_query.py         # NL query prompt construction
│   └── ocr_engine.py         # Tesseract wrapper + preprocessing
├── ui/
│   ├── sidebar.py            # Branding, connection, admin panel
│   ├── tab_reports.py        # Report Generator UI
│   ├── tab_query.py          # Data Query chat UI
│   └── tab_ocr.py            # Document OCR UI
├── utils/
│   ├── data_loader.py        # CSV/Excel file loading
│   └── formatters.py         # Currency formatting, metrics
├── sample_data/              # Synthetic test documents
├── scripts/                  # Generation utilities
├── docs/                     # Documentation & screenshots
├── .streamlit/
│   └── config.toml           # Dark theme configuration
└── requirements.txt""")

    add_heading_styled(doc, "4.2 Prerequisites", level=2)
    prereqs = [
        "Python 3.11+ (tested with 3.14)",
        "LM Studio with Gemma 4 model loaded (google/gemma-4-e4b)",
        "Tesseract OCR 5.x (optional but recommended): brew install tesseract",
        "pip packages: streamlit, openai, pandas, plotly, Pillow, pytesseract",
    ]
    for p in prereqs:
        add_bullet(doc, p)

    add_heading_styled(doc, "4.3 Installation & Setup", level=2)
    add_code_block(doc, """# 1. Clone the repository
git clone <repository-url>
cd FinVaultAI

# 2. Install Python dependencies
pip install -r requirements.txt

# 3. Install Tesseract OCR (macOS)
brew install tesseract

# 4. Start LM Studio and load google/gemma-4-e4b
# Ensure server is running on localhost:1234

# 5. Launch the application
streamlit run app.py""")

    add_heading_styled(doc, "4.4 Configuration", level=2)
    add_para(doc, (
        "All application configuration is centralized in config/settings.py. "
        "Key parameters include:"
    ))
    config_items = [
        ("LM_STUDIO_URL", "http://localhost:1234/v1 (default)"),
        ("DEFAULT_TEMPERATURE", "0.15 (low for deterministic financial output)"),
        ("DEFAULT_MAX_TOKENS", "2048"),
        ("DEFAULT_TOP_P", "0.95"),
        ("SUPPORTED_IMAGE_TYPES", "png, jpg, jpeg, tiff, bmp, webp"),
        ("DOC_TYPES", "Invoice, Work Order, Receipt, Purchase Order, Statement"),
    ]
    add_info_table(doc, config_items)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # 5. AGENT ENGINEERING & LLM CONFIGURATION
    # ════════════════════════════════════════════════════════════
    add_heading_styled(doc, "5. Agent Engineering & LLM Configuration", level=1)

    add_heading_styled(doc, "5.1 Prompt Engineering Strategy", level=2)
    add_para(doc, (
        "FinVault AI employs specialized system prompts that establish domain expertise "
        "for each module. Each prompt follows a structured pattern: role definition, "
        "behavioral constraints, output format specification, and safety guardrails."
    ))

    add_heading_styled(doc, "Report Generator Prompt Design", level=3)
    add_para(doc, (
        "The report system prompt establishes a 'Senior CPA / Financial Analyst' persona. "
        "It instructs the model to use Markdown formatting, include specific sections "
        "(Executive Summary, Key Findings, Detailed Analysis, Recommendations), "
        "cite exact figures from the data, and flag anomalies or risks."
    ))
    add_para(doc, "Key prompt elements:", bold=True)
    prompt_elements = [
        "Role: 'You are a senior CPA and financial analyst'",
        "Output structure: Enforced Markdown with H2/H3 headers",
        "Data context: Full CSV excerpt + pre-computed statistics",
        "Numeric precision: 'Always cite exact figures from the data provided'",
        "Anomaly detection: 'Highlight any anomalies, risks, or unusual patterns'",
    ]
    for pe in prompt_elements:
        add_bullet(doc, pe)

    add_heading_styled(doc, "Query Module Prompt Design", level=3)
    add_para(doc, (
        "The query prompt creates a 'financial data analyst' that receives the full "
        "data context (column names, types, row count, and CSV content). "
        "It is instructed to show calculations step-by-step, reference specific "
        "data points, and provide precise numerical answers."
    ))

    add_heading_styled(doc, "OCR Parsing Prompt Design", level=3)
    add_para(doc, (
        "Two prompts serve the OCR pipeline: (1) Vision prompt for direct image-to-text "
        "extraction when Tesseract is unavailable, and (2) Parse prompt that transforms "
        "raw OCR text into structured fields using document-type-specific templates "
        "(invoice fields differ from work order fields)."
    ))

    add_heading_styled(doc, "5.2 Inference Configuration Best Practices", level=2)
    config_best = [
        "Temperature 0.15: Low creativity ensures deterministic financial outputs — critical for audit consistency",
        "Max Tokens 2048: Sufficient for detailed reports while preventing runaway generation",
        "Top-P 0.95: Slight nucleus sampling allows natural language flow without hallucination risk",
        "System prompts are immutable: Users control temperature/tokens but cannot modify core prompts",
        "Context window management: Data is truncated to fit model context with priority on recent/relevant rows",
    ]
    for cb in config_best:
        add_bullet(doc, cb)

    add_heading_styled(doc, "5.3 LLM Client Architecture", level=2)
    add_para(doc, (
        "The LLM client (core/llm_client.py) wraps the OpenAI Python SDK, pointing it at "
        "the local LM Studio server. The LMStudioState dataclass tracks connection status, "
        "model name, and server URL. Key functions:"
    ))
    functions = [
        ("connect(url)", "Establishes connection to LM Studio, retrieves model list"),
        ("chat(state, system_prompt, user_prompt, ...)", "Standard text completion via /v1/chat/completions"),
        ("vision_chat(state, system_prompt, text, image)", "Multimodal vision request with base64-encoded image"),
    ]
    add_info_table(doc, functions)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # 6. SECURITY ARCHITECTURE
    # ════════════════════════════════════════════════════════════
    add_heading_styled(doc, "6. Security Architecture", level=1)

    add_heading_styled(doc, "6.1 Data Isolation Model", level=2)
    add_para(doc, (
        "FinVault AI implements a zero-trust data isolation model. Every data path is local:"
    ))
    isolation = [
        "Network: All API calls target localhost:1234 only. No DNS resolution, no outbound connections.",
        "Storage: No data is written to disk. All processing occurs in-memory within the Streamlit session.",
        "Session lifecycle: When the browser tab closes, all data is garbage-collected.",
        "Model weights: Gemma 4 runs entirely on-device via LM Studio. No cloud API fallback exists.",
        "OCR: Tesseract processes images locally. No cloud OCR services are invoked.",
    ]
    for item in isolation:
        add_bullet(doc, item)

    add_heading_styled(doc, "6.2 LLM Security Considerations", level=2)
    add_para(doc, (
        "Running an LLM locally eliminates the primary attack vector of cloud-based AI "
        "(data exfiltration via API calls), but introduces new considerations:"
    ))

    llm_security = [
        "Prompt injection defense: System prompts are hardcoded and not user-modifiable. User input is clearly delimited with markdown fences.",
        "Output sanitization: AI responses are rendered as Markdown. No raw HTML execution from model output.",
        "Token budget limits: Max tokens are capped to prevent resource exhaustion attacks.",
        "No persistent memory: The model has no access to previous sessions or user history beyond the current conversation.",
        "Model integrity: LM Studio loads verified model weights. No fine-tuning or weight modification occurs at runtime.",
    ]
    for ls in llm_security:
        add_bullet(doc, ls)

    add_heading_styled(doc, "6.3 Compliance Matrix", level=2)
    compliance = [
        ("SOX Compliance", "All financial data remains on-premises; audit trail via session logs"),
        ("GDPR Article 44", "No cross-border data transfer — processing is 100% local"),
        ("CCPA", "No personal data collection, storage, or sale"),
        ("HIPAA (if applicable)", "PHI never leaves the device; no BAA required"),
        ("Air-Gap Ready", "Operates without any network connectivity"),
    ]
    add_info_table(doc, compliance)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # 7. DEPLOYMENT GUIDE
    # ════════════════════════════════════════════════════════════
    add_heading_styled(doc, "7. Deployment Guide", level=1)

    add_heading_styled(doc, "7.1 System Requirements", level=2)
    requirements = [
        ("Operating System", "macOS 13+, Ubuntu 22.04+, Windows 11"),
        ("RAM", "16 GB minimum, 32 GB recommended"),
        ("GPU", "Optional but recommended (Metal on Mac, CUDA on Linux/Windows)"),
        ("Storage", "~8 GB for Gemma 4 model weights + 500 MB for application"),
        ("Python", "3.11 or later"),
        ("LM Studio", "Latest version with local server enabled"),
    ]
    add_info_table(doc, requirements)

    add_heading_styled(doc, "7.2 Deployment Process Flow", level=2)
    deploy_steps = [
        "Environment Setup: Install Python, LM Studio, Tesseract OCR",
        "Model Download: Load google/gemma-4-e4b in LM Studio (~8 GB)",
        "Server Configuration: Start LM Studio server on localhost:1234",
        "Application Install: Clone repo, pip install requirements",
        "Theme Configuration: .streamlit/config.toml sets dark theme defaults",
        "Validation: Run streamlit run app.py, verify connection status in sidebar",
        "Testing: Upload sample data, run each module (Reports, Query, OCR)",
        "Production: Configure auto-start scripts for LM Studio + Streamlit",
    ]
    for i, step in enumerate(deploy_steps, 1):
        add_bullet(doc, f"{i}. {step}")

    add_heading_styled(doc, "7.3 Health Checks", level=2)
    add_para(doc, "The Admin panel provides real-time system health monitoring:")
    checks = [
        "LLM Server: Green indicator when LM Studio is reachable at localhost:1234",
        "OCR Engine: Shows Tesseract version or warns if not installed (AI vision fallback activates)",
        "Model Info: Displays loaded model name, confirming Gemma 4 is active",
        "Inference Config: Shows current temperature, max tokens, and top-p values",
    ]
    for c in checks:
        add_bullet(doc, c)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # 8. USER DEMO GUIDE
    # ════════════════════════════════════════════════════════════
    add_heading_styled(doc, "8. User Demo Guide", level=1)

    add_para(doc, (
        "This section provides a step-by-step walkthrough for demonstrating FinVault AI "
        "to stakeholders. The demo uses sample data included in the repository."
    ), bold=True)

    add_heading_styled(doc, "8.1 Pre-Demo Checklist", level=2)
    checklist = [
        "LM Studio is running with Gemma 4 loaded and server active on port 1234",
        "Streamlit app is launched (streamlit run app.py)",
        "Browser is open to localhost:8501",
        "Sidebar shows green 'Connected' status with model name",
        "Sample data is available in sample_data/ directory",
    ]
    for c in checklist:
        add_bullet(doc, c)

    add_heading_styled(doc, "8.2 Demo Script: Report Generator (3 min)", level=2)
    demo_report = [
        "Click the 'Report Generator' tab",
        "Check 'Use sample data' → select sample_accounts.csv",
        "Point out: Data Preview table, column chips, quick metrics, and the bar chart",
        "Select report type: 'Executive Summary'",
        "Click 'Generate Report' — wait for AI analysis (~10-15 seconds)",
        "Highlight: Structured report with Executive Summary, Key Findings, Recommendations",
        "Show 'Download Report' button — emphasize no data left the machine",
    ]
    for i, step in enumerate(demo_report, 1):
        add_bullet(doc, f"Step {i}: {step}")

    add_heading_styled(doc, "8.3 Demo Script: Data Query (3 min)", level=2)
    demo_query = [
        "Switch to 'Data Query' tab",
        "Check 'Use sample data' → same CSV is loaded",
        "Point out column chips showing available fields",
        "Type: 'What is the total revenue for Q1 2026?'",
        "Show the AI response with step-by-step calculations",
        "Follow up: 'Which category has the highest expenses?'",
        "Emphasize: Context is maintained across questions",
    ]
    for i, step in enumerate(demo_query, 1):
        add_bullet(doc, f"Step {i}: {step}")

    add_heading_styled(doc, "8.4 Demo Script: Document OCR (3 min)", level=2)
    demo_ocr = [
        "Switch to 'Document OCR' tab",
        "Check 'Use sample docs' → select sample_invoice.png",
        "Click 'Extract & Parse'",
        "Show: OCR confidence score, word count, raw extracted text",
        "Scroll to: AI-parsed structured data (invoice number, date, amounts)",
        "Highlight: 'Download Raw Text' and 'Download Parsed Data' buttons",
        "Optional: Toggle 'Pre-process' checkbox for low-quality scans",
    ]
    for i, step in enumerate(demo_ocr, 1):
        add_bullet(doc, f"Step {i}: {step}")

    add_heading_styled(doc, "8.5 Key Talking Points", level=2)
    talking = [
        "Privacy: 'Notice the URL — localhost:8501. Everything runs on this machine.'",
        "Speed: 'Report generated in under 15 seconds — no API latency.'",
        "Compliance: 'Zero cloud means SOX/GDPR compliance is built into the architecture.'",
        "Cost: 'No API fees. No subscriptions. One-time setup.'",
        "Flexibility: 'The model handles financial jargon, calculates metrics, and parses any document type.'",
    ]
    for t in talking:
        add_bullet(doc, t)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # 9. PRODUCT ROADMAP
    # ════════════════════════════════════════════════════════════
    add_heading_styled(doc, "9. Product Roadmap", level=1)

    add_screenshot(doc, "06_roadmap.png", "Figure 6: Product roadmap — phased delivery through Q1 2027")

    add_heading_styled(doc, "Q2 2026 — Current Release (v3.0)", level=2)
    q2 = [
        "Core platform: Report Generator, Data Query, Document OCR",
        "Google Gemma 4 integration via LM Studio",
        "Tesseract 5.5 OCR with vision fallback",
        "Dark theme UI with Admin panel",
        "Sample data and synthetic test documents",
    ]
    for item in q2:
        add_bullet(doc, item)

    add_heading_styled(doc, "Q3 2026 — Intelligence Layer", level=2)
    q3 = [
        "RAG (Retrieval-Augmented Generation) pipeline for document-grounded answers",
        "Anomaly detection engine with configurable thresholds",
        "Multi-document cross-referencing (e.g., match invoices to purchase orders)",
        "Batch processing mode for bulk document OCR",
        "Export to PDF format for reports",
    ]
    for item in q3:
        add_bullet(doc, item)

    add_heading_styled(doc, "Q4 2026 — Enterprise Features", level=2)
    q4 = [
        "Multi-user support with role-based access control (RBAC)",
        "Audit logging and compliance reporting",
        "API endpoints for integration with existing ERP systems",
        "Custom report templates with branding options",
        "Scheduled report generation (cron-based)",
    ]
    for item in q4:
        add_bullet(doc, item)

    add_heading_styled(doc, "Q1 2027 — Scale & Observability", level=2)
    q1_27 = [
        "Model performance monitoring and inference analytics",
        "A/B testing framework for prompt optimization",
        "Advanced agentic workflows: multi-step financial analysis pipelines",
        "Kubernetes deployment option for team environments",
        "Telemetry dashboard (local-only) for usage metrics",
    ]
    for item in q1_27:
        add_bullet(doc, item)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # 10. APPENDIX
    # ════════════════════════════════════════════════════════════
    add_heading_styled(doc, "10. Appendix", level=1)

    add_heading_styled(doc, "A. API Reference", level=2)
    add_para(doc, "LM Studio exposes an OpenAI-compatible API at localhost:1234:")
    api_endpoints = [
        ("GET /v1/models", "List available models"),
        ("POST /v1/chat/completions", "Text/vision chat completion"),
        ("GET /v1/models/{id}", "Model metadata"),
    ]
    add_info_table(doc, api_endpoints)

    add_heading_styled(doc, "B. Supported Document Types", level=2)
    doc_types = [
        ("Invoice", "Invoice #, Date, Vendor, Line Items, Subtotal, Tax, Total"),
        ("Work Order", "WO #, Date, Client, Description, Labor, Materials, Total"),
        ("Receipt", "Store, Date, Items, Subtotal, Tax, Total, Payment Method"),
        ("Purchase Order", "PO #, Vendor, Items, Quantities, Unit Prices, Total"),
        ("Statement", "Account, Period, Opening Balance, Transactions, Closing Balance"),
    ]
    add_info_table(doc, doc_types)

    add_heading_styled(doc, "C. Troubleshooting", level=2)
    troubles = [
        "Connection failed: Ensure LM Studio server is running on port 1234. Check sidebar status.",
        "Low OCR confidence: Enable 'Pre-process' checkbox. Use 300+ DPI scans for best results.",
        "Slow inference: Ensure GPU acceleration is enabled in LM Studio preferences.",
        "Module not found: Run pip install -r requirements.txt to install all dependencies.",
        "Tesseract not detected: Install via 'brew install tesseract' (macOS) or 'apt install tesseract-ocr' (Ubuntu).",
    ]
    for t in troubles:
        add_bullet(doc, t)

    # ── Footer note ──
    add_para(doc, "", size=6)
    add_para(doc, (
        "FinVault AI v3.0 — Confidential Technical Documentation — April 2026\n"
        "All processing runs locally. No data is transmitted. "
        "Powered by Google Gemma 4 via LM Studio."
    ), size=9, italic=True, color=RGBColor(0x94, 0xA3, 0xB8), align="center")

    # ── Save ──
    doc.save(str(OUTPUT))
    print(f"DOCX created: {OUTPUT}")


if __name__ == "__main__":
    build_document()
