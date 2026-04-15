#!/usr/bin/env python3
"""
Generate FinVault AI Technical Guide (.docx)
Uses python-docx to create a professionally formatted Word document.
"""
import sys
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    import os
except ImportError:
    print("python-docx not installed. Installing...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    import os

# --- Constants ---
INDIGO = RGBColor(0x63, 0x66, 0xF1)
DARK_INDIGO = RGBColor(0x31, 0x33, 0x78)
DARKER_INDIGO = RGBColor(0x1E, 0x1B, 0x4B)
LIGHT_INDIGO = RGBColor(0xA5, 0xB4, 0xFC)
NEAR_BLACK = RGBColor(0x1F, 0x20, 0x37)
DARK_GRAY = RGBColor(0x37, 0x41, 0x51)
MED_GRAY = RGBColor(0x6B, 0x72, 0x80)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG = RGBColor(0xEE, 0xF2, 0xFF)  # indigo-50
CODE_BG = "F1F5F9"  # slate-100

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
SCREENSHOTS_DIR = os.path.join(BASE_DIR, "screenshots")
OUTPUT_PATH = os.path.join(BASE_DIR, "FinVault_AI_Technical_Guide.docx")

doc = Document()

# =========================================================================
# STYLES
# =========================================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)
font.color.rgb = NEAR_BLACK
pf = style.paragraph_format
pf.space_after = Pt(6)
pf.space_before = Pt(0)
pf.line_spacing = 1.15

# --- Heading styles ---
for level, (size, color) in enumerate([
    (Pt(22), DARK_INDIGO),   # Heading 1
    (Pt(16), DARK_INDIGO),   # Heading 2
    (Pt(13), INDIGO),        # Heading 3
], start=1):
    h_style = doc.styles[f'Heading {level}']
    h_style.font.name = 'Calibri'
    h_style.font.size = size
    h_style.font.color.rgb = color
    h_style.font.bold = True
    h_style.paragraph_format.space_before = Pt(18 if level == 1 else 14 if level == 2 else 10)
    h_style.paragraph_format.space_after = Pt(8)
    if level == 1:
        # Bottom border for H1
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="8" w:space="4" w:color="6366F1"/>'
            f'</w:pBdr>'
        )
        h_style.element.pPr.append(pBdr)

# Title style
title_style = doc.styles['Title']
title_style.font.name = 'Calibri'
title_style.font.size = Pt(36)
title_style.font.color.rgb = DARK_INDIGO
title_style.font.bold = True

# Subtitle style
subtitle_style = doc.styles['Subtitle']
subtitle_style.font.name = 'Calibri'
subtitle_style.font.size = Pt(16)
subtitle_style.font.color.rgb = INDIGO
subtitle_style.font.bold = False
subtitle_style.font.italic = True

# List Bullet style
try:
    list_style = doc.styles['List Bullet']
    list_style.font.name = 'Calibri'
    list_style.font.size = Pt(10.5)
    list_style.font.color.rgb = NEAR_BLACK
except:
    pass

# =========================================================================
# HELPER FUNCTIONS
# =========================================================================

def add_page_break():
    doc.add_page_break()

def add_heading_1(text):
    return doc.add_heading(text, level=1)

def add_heading_2(text):
    return doc.add_heading(text, level=2)

def add_heading_3(text):
    return doc.add_heading(text, level=3)

def add_para(text, bold=False, italic=False, size=None, color=None, alignment=None, space_after=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if size:
        run.font.size = size
    if color:
        run.font.color.rgb = color
    if alignment:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = space_after
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.5 * (level + 1))
    return p

def add_bold_bullet(bold_text, normal_text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run_b = p.add_run(bold_text)
    run_b.bold = True
    p.add_run(normal_text)
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.5 * (level + 1))
    return p

def add_code_block(text):
    """Add a shaded code block."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # Add shading
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{CODE_BG}" w:val="clear"/>')
    p.paragraph_format.element.get_or_add_pPr().append(shading)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = DARK_GRAY
    return p

def add_screenshot(filename, caption, width=Inches(5.8)):
    """Add an image with a caption."""
    filepath = os.path.join(SCREENSHOTS_DIR, filename)
    if os.path.exists(filepath):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(filepath, width=width)
        # Caption
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_cap = cap.add_run(caption)
        run_cap.font.size = Pt(9)
        run_cap.font.italic = True
        run_cap.font.color.rgb = MED_GRAY
        cap.paragraph_format.space_after = Pt(12)
    else:
        add_para(f"[Image not found: {filename}]", italic=True, color=MED_GRAY)

def add_accent_line():
    """Add a thin indigo horizontal rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="6366F1"/>'
        f'</w:pBdr>'
    )
    p.paragraph_format.element.get_or_add_pPr().append(pBdr)

def add_table_simple(headers, rows, col_widths=None):
    """Add a styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = WHITE
        run.font.name = 'Calibri'
        # Indigo background
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="4F46E5" w:val="clear"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = "F5F3FF" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9.5)
            run.font.name = 'Calibri'
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg}" w:val="clear"/>')
            cell._tc.get_or_add_tcPr().append(shading)

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = w

    doc.add_paragraph()  # spacing
    return table


# =========================================================================
# PAGE SETUP
# =========================================================================
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.8)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)

# Header
header = section.header
header.is_linked_to_previous = False
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run_h = hp.add_run("FinVault AI  |  Technical Guide v3.0")
run_h.font.size = Pt(8)
run_h.font.color.rgb = MED_GRAY
run_h.font.name = 'Calibri'
# Header bottom border
pBdr = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'  <w:bottom w:val="single" w:sz="4" w:space="4" w:color="C7D2FE"/>'
    f'</w:pBdr>'
)
hp._element.get_or_add_pPr().append(pBdr)

# Footer
footer = section.footer
footer.is_linked_to_previous = False
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_f = fp.add_run("Confidential \u2014 Internal Use Only")
run_f.font.size = Pt(8)
run_f.font.color.rgb = MED_GRAY
run_f.font.name = 'Calibri'


# =========================================================================
# COVER PAGE
# =========================================================================
for _ in range(6):
    doc.add_paragraph()

# Accent line
add_accent_line()

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("FinVault AI")
run.font.name = 'Calibri'
run.font.size = Pt(42)
run.font.color.rgb = DARK_INDIGO
run.bold = True
p.paragraph_format.space_after = Pt(4)

# Subtitle
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Product Design, Development & Deployment Guide")
run.font.name = 'Calibri'
run.font.size = Pt(16)
run.font.color.rgb = INDIGO
run.italic = True
p.paragraph_format.space_after = Pt(4)

# Accent line
add_accent_line()

# Version and date
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
run = p.add_run("Version 3.0  |  April 2026")
run.font.size = Pt(11)
run.font.color.rgb = MED_GRAY

# Tagline
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(30)
run = p.add_run("On-device financial intelligence. Private by design.")
run.font.size = Pt(12)
run.font.color.rgb = INDIGO
run.italic = True

for _ in range(4):
    doc.add_paragraph()

# Confidentiality
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CONFIDENTIAL \u2014 INTERNAL USE ONLY")
run.font.size = Pt(9)
run.font.color.rgb = MED_GRAY
run.bold = True
run.font.all_caps = True

add_page_break()


# =========================================================================
# TABLE OF CONTENTS
# =========================================================================
add_heading_1("Table of Contents")
add_para("")  # space

toc_items = [
    ("1.", "Executive Overview", "3"),
    ("2.", "Product Design", "4"),
    ("", "2.1  Design Philosophy", "4"),
    ("", "2.2  User Personas", "4"),
    ("", "2.3  Feature Set", "5"),
    ("", "2.4  UX Design Decisions", "6"),
    ("3.", "System Architecture", "7"),
    ("", "3.1  High-Level Architecture", "7"),
    ("", "3.2  Component Breakdown", "7"),
    ("", "3.3  Data Flow Diagram", "8"),
    ("4.", "Agent Engineering & LLM Configuration", "9"),
    ("", "4.1  Secure LLM Usage", "9"),
    ("", "4.2  Prompt Engineering Best Practices", "9"),
    ("", "4.3  Configuration Parameters", "10"),
    ("", "4.4  Multi-Modal Capabilities", "10"),
    ("5.", "Development Technical Guide", "11"),
    ("", "5.1  Project Structure", "11"),
    ("", "5.2  Technology Stack", "12"),
    ("", "5.3  Development Setup", "12"),
    ("", "5.4  Code Quality Best Practices", "13"),
    ("6.", "Deployment Process", "14"),
    ("", "6.1  Local Deployment", "14"),
    ("", "6.2  Docker Deployment", "15"),
    ("", "6.3  Enterprise Deployment", "15"),
    ("7.", "User Demo Guide", "16"),
    ("", "7.1  Report Generation Demo", "16"),
    ("", "7.2  Natural Language Query Demo", "17"),
    ("", "7.3  OCR Document Processing Demo", "18"),
    ("8.", "Roadmap \u2014 Future Enhancements", "19"),
    ("", "8.1  Phase 2 \u2014 Intelligence (Q3 2026)", "19"),
    ("", "8.2  Phase 3 \u2014 Enterprise (Q4 2026)", "19"),
    ("", "8.3  Phase 4 \u2014 Scale (Q1 2027)", "20"),
    ("", "8.4  Security Enhancements", "20"),
    ("", "8.5  Optimization", "20"),
    ("9.", "Appendix", "21"),
    ("", "9.1  System Prompts Reference", "21"),
    ("", "9.2  Supported File Types", "22"),
    ("", "9.3  Sample Data Files", "22"),
]

for num, title, page in toc_items:
    p = doc.add_paragraph()
    is_main = num != ""
    indent = Inches(0) if is_main else Inches(0.4)
    p.paragraph_format.left_indent = indent
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2 if not is_main else 6)

    if is_main:
        run_num = p.add_run(f"{num}  ")
        run_num.bold = True
        run_num.font.size = Pt(11)
        run_num.font.color.rgb = INDIGO
    run_title = p.add_run(title)
    run_title.bold = is_main
    run_title.font.size = Pt(11 if is_main else 10)
    run_title.font.color.rgb = DARK_INDIGO if is_main else DARK_GRAY

    # Dot leader + page number
    run_dots = p.add_run(f"  {'.' * (60 if is_main else 50)}  ")
    run_dots.font.size = Pt(8)
    run_dots.font.color.rgb = RGBColor(0xC7, 0xD2, 0xFE)
    run_pg = p.add_run(page)
    run_pg.font.size = Pt(10)
    run_pg.font.color.rgb = MED_GRAY

add_page_break()


# =========================================================================
# 1. EXECUTIVE OVERVIEW
# =========================================================================
add_heading_1("1. Executive Overview")

add_para(
    "FinVault AI is a secure, 100% local financial intelligence platform designed to bring the power of "
    "modern large language models to financial professionals without compromising data privacy or security. "
    "Built on a zero-cloud, air-gapped architecture, the application processes all data on-device using "
    "Google Gemma 4 running through LM Studio's local inference server."
)

add_para(
    "The platform eliminates the fundamental tension between AI capability and data security that plagues "
    "cloud-based solutions. No financial data, prompts, or inference results ever leave the user's machine. "
    "This makes FinVault AI suitable for environments with the strictest regulatory and compliance requirements, "
    "including SOX-governed organizations, auditing firms, and financial institutions."
)

add_heading_3("Target Users")
add_bullet("Financial professionals (CFOs, Controllers, Financial Analysts)")
add_bullet("Accountants and bookkeepers managing ledger data")
add_bullet("Auditors requiring secure document analysis")
add_bullet("Accounts Payable / Accounts Receivable clerks processing invoices and receipts")

add_heading_3("Core Capabilities")
add_table_simple(
    ["Capability", "Description"],
    [
        ["Report Generation", "AI-generated financial reports from CSV/Excel data \u2014 7 report types including Executive Summary, Ledger Analysis, Expense Breakdown, and more"],
        ["Natural Language Query", "Chat-style data analysis: ask questions about your financial data in plain English and receive data-grounded answers with step-by-step calculations"],
        ["Document OCR", "Extract and parse structured data from invoices, work orders, receipts, and purchase orders using Tesseract OCR + Gemma AI parsing"],
        ["Admin Panel", "Configure model parameters, monitor system health, and manage inference settings in real time"],
    ],
    col_widths=[Inches(1.8), Inches(4.6)]
)

add_heading_3("Key Differentiator")
add_para(
    "FinVault AI's core differentiator is its zero-cloud, air-gapped architecture. The entire AI inference "
    "pipeline runs locally on the user's hardware. There are no API keys to manage, no cloud endpoints to "
    "secure, and no data egress to monitor. The application can operate on a completely disconnected network, "
    "making it ideal for sensitive financial environments."
)

add_page_break()


# =========================================================================
# 2. PRODUCT DESIGN
# =========================================================================
add_heading_1("2. Product Design")

# 2.1
add_heading_2("2.1 Design Philosophy")
add_para(
    "FinVault AI is built on three foundational design principles that guide every architectural and UX decision:"
)
add_bold_bullet("Privacy-First: ", "All data processing occurs locally. No information is transmitted to external servers. The application requires no internet connectivity to function after initial setup.")
add_bold_bullet("Local-First: ", "The complete AI inference pipeline runs on the user's hardware via LM Studio. This eliminates latency from network round-trips and ensures availability regardless of cloud service status.")
add_bold_bullet("Modular Architecture: ", "The codebase follows strict separation of concerns \u2014 UI components, core logic, configuration, and utilities are independently maintained modules that can be extended or replaced without affecting the rest of the system.")

# 2.2
add_heading_2("2.2 User Personas")

add_table_simple(
    ["Persona", "Role", "Primary Use Case", "Key Need"],
    [
        ["CFO / Controller", "Executive oversight", "Executive summaries, trend analysis, variance reports", "Quick, accurate financial insights from raw data"],
        ["Staff Accountant", "Day-to-day accounting", "Ledger analysis, data queries, expense tracking", "Efficient data exploration without manual spreadsheet work"],
        ["Auditor", "Compliance review", "Document verification, cross-referencing, anomaly detection", "Secure processing of sensitive financial documents"],
        ["AP/AR Clerk", "Transaction processing", "Invoice OCR, receipt parsing, PO extraction", "Fast, accurate extraction of structured data from documents"],
    ],
    col_widths=[Inches(1.2), Inches(1.2), Inches(2.2), Inches(1.8)]
)

# 2.3
add_heading_2("2.3 Feature Set")

add_heading_3("Report Generator")
add_para("The Report Generator produces AI-authored financial reports from uploaded CSV or Excel data. Seven report types are available:")
for rt in ["Executive Summary", "Detailed Ledger Analysis", "Expense Breakdown", "Revenue vs. Expense Comparison", "Cash Flow Summary", "Monthly Trend Analysis", "Budget Variance Report"]:
    add_bullet(rt)
add_para(
    "Each report includes automated statistical pre-computation (totals, means, category breakdowns) "
    "that is injected into the LLM prompt alongside the raw data, ensuring reports are grounded in actual figures."
)

add_heading_3("Natural Language Query")
add_para(
    "The Data Query module enables conversational interaction with financial data. Users type questions "
    "in plain English (e.g., 'What is the total revenue for Q1 2026?') and receive data-grounded answers "
    "with step-by-step calculations. Chat history is maintained via Streamlit session state, allowing "
    "multi-turn conversations with context carryover."
)

add_heading_3("Document OCR")
add_para(
    "The OCR pipeline processes financial documents in two stages: (1) Tesseract OCR extracts raw text "
    "from images with confidence scoring, and (2) Gemma AI parses the extracted text into structured "
    "fields using document-type-specific templates. Supported document types include Invoices, Work Orders, "
    "Receipts, Purchase Orders, and General Documents. When Tesseract is unavailable, the system falls "
    "back to Gemma 4's vision capabilities for direct image-to-text extraction."
)

add_heading_3("Admin Panel")
add_para(
    "The Admin Panel provides real-time control over the AI inference pipeline. Users can adjust model "
    "temperature, maximum tokens, Top-P sampling, and other parameters. The panel also displays system "
    "health information including LM Studio connection status, loaded model details, and resource utilization."
)

# 2.4
add_heading_2("2.4 UX Design Decisions")
add_bold_bullet("Dark Theme: ", "Optimized for extended professional sessions. Reduces eye strain during prolonged financial analysis work.")
add_bold_bullet("Tab-Based Navigation: ", "Primary features (Reports, Query, OCR) are organized as tabs for quick switching without page reloads.")
add_bold_bullet("Progressive Disclosure: ", "Advanced options and detailed results are housed within Streamlit expanders, keeping the default view clean and focused.")
add_bold_bullet("Responsive Layout: ", "The interface adapts to different screen sizes using Streamlit's column system, ensuring usability on both large monitors and laptops.")
add_bold_bullet("Real-Time Feedback: ", "Spinners and status messages provide immediate feedback during AI inference, OCR processing, and data loading operations.")

add_screenshot("01_dashboard_report.png", "Figure 1: Report Generator \u2014 data preview, metrics, and chart visualization")

add_page_break()


# =========================================================================
# 3. SYSTEM ARCHITECTURE
# =========================================================================
add_heading_1("3. System Architecture")

# 3.1
add_heading_2("3.1 High-Level Architecture")
add_para(
    "FinVault AI follows a strictly local client-server architecture where all components run on the "
    "user's machine within an air-gapped boundary:"
)

add_code_block(
    "User Browser\n"
    "    \u2502\n"
    "    \u25BC\n"
    "Streamlit Web App (localhost:8501)\n"
    "    \u2502                    \u2502\n"
    "    \u25BC                    \u25BC\n"
    "LM Studio API           Tesseract OCR\n"
    "(localhost:1234)         (local binary)\n"
    "    \u2502\n"
    "    \u25BC\n"
    "Google Gemma 4\n"
    "(on-device model)"
)

add_para(
    "The Streamlit application serves as the central orchestrator. It receives user input through the browser, "
    "routes requests to either the LM Studio API (for LLM inference) or Tesseract (for OCR), and renders "
    "results back to the user. All communication occurs over localhost \u2014 no external network requests are made."
)

# 3.2
add_heading_2("3.2 Component Breakdown")

add_table_simple(
    ["File / Module", "Purpose"],
    [
        ["app.py", "Entry point. Page configuration, global CSS injection, tab routing, session state initialization."],
        ["config/settings.py", "Centralized constants: app identity, LM Studio defaults, LLM parameters, system prompts, OCR config, templates."],
        ["core/llm_client.py", "OpenAI-compatible client for LM Studio. Handles text completion and vision (multi-modal) API calls with error handling."],
        ["core/ocr_engine.py", "Tesseract wrapper with image preprocessing. Extracts text with confidence scoring, handles fallback to vision API."],
        ["core/report_generator.py", "Prompt engineering module for financial reports. Builds context-rich prompts with data + statistics + report-type instructions."],
        ["core/data_query.py", "Natural language question-to-answer pipeline. Injects CSV data + summary into prompts, manages chat context."],
        ["ui/sidebar.py", "Sidebar component: data upload, sample data toggle, file type selection, data preview."],
        ["ui/tab_reports.py", "Report Generator tab: metrics display, chart visualization, report type selection, AI report generation."],
        ["ui/tab_query.py", "Data Query tab: chat interface, message history, query input, AI response rendering."],
        ["ui/tab_ocr.py", "Document OCR tab: image upload, document type selection, OCR extraction, AI parsing, download options."],
        ["utils/data_loader.py", "File loading utilities for CSV, XLSX, XLS formats with error handling and validation."],
        ["utils/formatters.py", "Output formatting helpers: monetary values, markdown rendering, data summarization."],
    ],
    col_widths=[Inches(1.8), Inches(4.6)]
)

# 3.3
add_heading_2("3.3 Data Flow Diagram")
add_para(
    "The following diagram illustrates the complete system architecture, showing how all components "
    "interact within the local air-gapped boundary:"
)

add_screenshot("05_architecture.png", "Figure 2: System Architecture \u2014 all components run locally within an air-gapped boundary")

add_page_break()


# =========================================================================
# 4. AGENT ENGINEERING & LLM CONFIGURATION
# =========================================================================
add_heading_1("4. Agent Engineering & LLM Configuration")

# 4.1
add_heading_2("4.1 How the LLM is Used Securely")
add_para(
    "FinVault AI treats security as a non-negotiable architectural constraint. The LLM integration "
    "is designed to ensure that no financial data ever leaves the user's machine:"
)
add_bold_bullet("Local Inference Only: ", "All LLM calls go to localhost:1234 via the OpenAI-compatible API provided by LM Studio. There is no fallback to cloud endpoints.")
add_bold_bullet("No External API Keys: ", "The only API key used is the local placeholder 'lm-studio', which never leaves the machine. No external service credentials are required.")
add_bold_bullet("In-Memory Data Processing: ", "Financial data exists only in Streamlit session state during the active session. No data is persisted to disk by the application.")
add_bold_bullet("Zero Telemetry: ", "No usage analytics, error reporting, or telemetry data is collected or transmitted. The application operates in complete isolation.")
add_bold_bullet("User-Owned Hardware: ", "The Gemma 4 model runs on the user's own GPU or CPU. Model weights are stored locally and managed through LM Studio.")

# 4.2
add_heading_2("4.2 Prompt Engineering Best Practices")
add_para("FinVault AI employs several prompt engineering techniques to ensure accurate, reliable financial analysis:")

add_heading_3("Expert Persona System Prompts")
add_para(
    "Each module uses a specialized system prompt that defines the AI's role and behavioral constraints. "
    "The Report Generator uses a CPA/financial analyst persona, the Data Query module uses a financial "
    "data analyst persona, and the OCR module uses a document parser persona."
)

add_heading_3("Context Injection Strategy")
add_para(
    "Prompts are enriched with comprehensive context to ground the model's responses in actual data:"
)
add_bullet("Full CSV data (up to 8,000 characters) is injected directly into the prompt")
add_bullet("Pre-computed statistical summaries (totals, means, counts, category breakdowns) are included")
add_bullet("Aggregated metrics are calculated before prompt construction to reduce hallucination risk")

add_heading_3("Structured Output Instructions")
add_para("System prompts specify exact output formatting requirements:")
add_bullet("Markdown tables for tabular data presentation")
add_bullet("Defined report sections (Executive Summary, Key Metrics, Analysis, Risks, Recommendations)")
add_bullet("Monetary value formatting ($X,XXX.XX with commas and two decimal places)")
add_bullet("Step-by-step calculation display for query responses")

add_heading_3("Temperature and Guard Rails")
add_bold_bullet("Temperature: ", "Set to 0.15 by default for financial work, prioritizing deterministic, consistent outputs over creative variation. Adjustable via the Admin Panel.")
add_bold_bullet("Data Grounding: ", "Prompts include explicit instructions to 'answer ONLY from data provided' and to return 'N/A' when information is not available in the dataset.")
add_bold_bullet("Anomaly Flagging: ", "The query prompt instructs the model to proactively flag suspicious or inconsistent values encountered during analysis.")

# 4.3
add_heading_2("4.3 Configuration Parameters")
add_para("The following inference parameters are configurable through the Admin Panel:")

add_table_simple(
    ["Parameter", "Range", "Default", "Purpose"],
    [
        ["Temperature", "0.0 \u2013 1.0", "0.15", "Controls output randomness. Lower values produce more deterministic results."],
        ["Max Tokens", "256 \u2013 4,096", "2,048 (reports: 3,072)", "Maximum length of generated response."],
        ["Top-P", "0.0 \u2013 1.0", "0.95", "Nucleus sampling threshold. Controls diversity of token selection."],
        ["Frequency Penalty", "0.0 \u2013 2.0", "0.0", "Penalizes repeated tokens to reduce redundancy."],
        ["Presence Penalty", "0.0 \u2013 2.0", "0.0", "Encourages the model to discuss new topics."],
        ["Model", "Any loaded model", "Gemma 4", "Supports any model loaded in LM Studio."],
        ["Endpoint URL", "Any localhost URL", "http://localhost:1234/v1", "LM Studio API endpoint."],
    ],
    col_widths=[Inches(1.3), Inches(1.1), Inches(1.5), Inches(2.5)]
)

# 4.4
add_heading_2("4.4 Multi-Modal Capabilities")
add_para(
    "FinVault AI leverages Gemma 4's vision capabilities as a fallback mechanism for OCR processing. "
    "When Tesseract is unavailable or produces low-confidence results, the application encodes document "
    "images as Base64 and sends them to the vision-capable endpoint of the LM Studio API. This enables "
    "direct image-to-text extraction without external OCR dependencies."
)
add_para(
    "The vision pipeline uses the same OpenAI-compatible API format, sending images as base64-encoded "
    "data URLs within the message content array. This approach maintains the zero-cloud architecture "
    "while providing robust document processing capabilities."
)

add_screenshot("04_admin_panel.png", "Figure 3: Admin Panel \u2014 model, inference settings, and system health monitoring")

add_page_break()


# =========================================================================
# 5. DEVELOPMENT TECHNICAL GUIDE
# =========================================================================
add_heading_1("5. Development Technical Guide")

# 5.1
add_heading_2("5.1 Project Structure")
add_para("The application follows a modular directory structure with clear separation of concerns:")

add_code_block(
    "LocalGemmaApp/\n"
    "\u251C\u2500\u2500 app.py                     # Entry point, page config, tab routing\n"
    "\u251C\u2500\u2500 requirements.txt            # Python dependencies\n"
    "\u251C\u2500\u2500 .streamlit/\n"
    "\u2502   \u2514\u2500\u2500 config.toml            # Streamlit server & theme config\n"
    "\u251C\u2500\u2500 config/\n"
    "\u2502   \u2514\u2500\u2500 settings.py            # All constants, prompts, templates\n"
    "\u251C\u2500\u2500 core/\n"
    "\u2502   \u251C\u2500\u2500 __init__.py\n"
    "\u2502   \u251C\u2500\u2500 llm_client.py          # OpenAI-compatible LM Studio client\n"
    "\u2502   \u251C\u2500\u2500 ocr_engine.py          # Tesseract wrapper + preprocessing\n"
    "\u2502   \u251C\u2500\u2500 report_generator.py    # Prompt engineering for reports\n"
    "\u2502   \u2514\u2500\u2500 data_query.py          # NL question \u2192 data answer\n"
    "\u251C\u2500\u2500 ui/\n"
    "\u2502   \u251C\u2500\u2500 __init__.py\n"
    "\u2502   \u251C\u2500\u2500 sidebar.py             # Data upload & preview sidebar\n"
    "\u2502   \u251C\u2500\u2500 tab_reports.py         # Report Generator tab\n"
    "\u2502   \u251C\u2500\u2500 tab_query.py           # Data Query chat tab\n"
    "\u2502   \u2514\u2500\u2500 tab_ocr.py             # Document OCR tab\n"
    "\u251C\u2500\u2500 utils/\n"
    "\u2502   \u251C\u2500\u2500 __init__.py\n"
    "\u2502   \u251C\u2500\u2500 data_loader.py         # CSV/XLSX/XLS file loading\n"
    "\u2502   \u2514\u2500\u2500 formatters.py          # Output formatting helpers\n"
    "\u251C\u2500\u2500 sample_data/\n"
    "\u2502   \u251C\u2500\u2500 sample_accounts.csv\n"
    "\u2502   \u251C\u2500\u2500 sample_accounts_payable.csv\n"
    "\u2502   \u251C\u2500\u2500 sample_invoice.png\n"
    "\u2502   \u251C\u2500\u2500 sample_invoice.pdf\n"
    "\u2502   \u251C\u2500\u2500 sample_purchase_order.pdf\n"
    "\u2502   \u251C\u2500\u2500 sample_receipt.png\n"
    "\u2502   \u2514\u2500\u2500 sample_workorder.png\n"
    "\u251C\u2500\u2500 scripts/\n"
    "\u2502   \u251C\u2500\u2500 generate_samples.py    # Synthetic test data generator\n"
    "\u2502   \u251C\u2500\u2500 generate_screenshots.py\n"
    "\u2502   \u2514\u2500\u2500 launcher.py            # Application launcher script\n"
    "\u251C\u2500\u2500 assets/                        # Static assets (icons, CSS)\n"
    "\u2514\u2500\u2500 docs/\n"
    "    \u251C\u2500\u2500 screenshots/               # Application screenshots\n"
    "    \u2514\u2500\u2500 FinVault_AI_Technical_Guide.docx"
)

# 5.2
add_heading_2("5.2 Technology Stack")

add_table_simple(
    ["Technology", "Version", "Purpose"],
    [
        ["Python", "3.9+", "Core runtime"],
        ["Streamlit", "1.30+", "Web UI framework \u2014 reactive, Python-native"],
        ["OpenAI SDK", "1.10+", "Client for OpenAI-compatible API (LM Studio)"],
        ["Pandas", "2.0+", "Data loading, manipulation, and statistical computation"],
        ["Plotly", "5.18+", "Interactive chart visualization"],
        ["Pillow", "10.0+", "Image processing and format handling"],
        ["pytesseract", "0.3.10+", "Python wrapper for Tesseract OCR engine"],
        ["openpyxl", "3.1+", "Excel file (.xlsx) reading support"],
        ["LM Studio", "Latest", "Local LLM inference server with OpenAI-compatible API"],
        ["Google Gemma 4", "Latest", "On-device large language model (text + vision)"],
        ["Tesseract OCR", "5.5+", "Open-source OCR engine (Homebrew / apt)"],
    ],
    col_widths=[Inches(1.5), Inches(1.0), Inches(3.9)]
)

# 5.3
add_heading_2("5.3 Development Setup")
add_para("Follow these steps to set up a local development environment:")

add_heading_3("Step 1: Clone the Repository")
add_code_block("git clone <repository-url>\ncd LocalGemmaApp")

add_heading_3("Step 2: Install Python Dependencies")
add_code_block("pip install -r requirements.txt")
add_para("This installs: streamlit, openai, pandas, openpyxl, Pillow, pytesseract, plotly")

add_heading_3("Step 3: Install Tesseract OCR")
add_code_block(
    "# macOS (Homebrew)\n"
    "brew install tesseract\n"
    "\n"
    "# Ubuntu / Debian\n"
    "sudo apt-get install tesseract-ocr\n"
    "\n"
    "# Verify installation\n"
    "tesseract --version"
)

add_heading_3("Step 4: Configure LM Studio")
add_bullet("Download and install LM Studio from https://lmstudio.ai")
add_bullet("Download Google Gemma 4 model (or any compatible model)")
add_bullet("Start the local server (default: localhost:1234)")
add_bullet("Ensure the server is running before launching FinVault AI")

add_heading_3("Step 5: Launch the Application")
add_code_block("streamlit run app.py")
add_para("The application will be available at http://localhost:8501 in your browser.")

# 5.4
add_heading_2("5.4 Code Quality Best Practices")
add_bold_bullet("Modular Architecture: ", "Strict separation of UI components (ui/), core logic (core/), configuration (config/), and utilities (utils/). Each module has a single responsibility.")
add_bold_bullet("Type Hints: ", "Python type annotations are used throughout the codebase to improve readability and enable static analysis.")
add_bold_bullet("Centralized Configuration: ", "All constants, prompts, templates, and default values live in config/settings.py \u2014 the single source of truth.")
add_bold_bullet("Error Handling: ", "Graceful degradation throughout. OCR falls back to vision API when Tesseract is unavailable. LLM connection failures display informative error messages.")
add_bold_bullet("Session State Management: ", "Streamlit session state is used to persist chat history, loaded data, and user preferences across reruns without server-side storage.")

add_page_break()


# =========================================================================
# 6. DEPLOYMENT PROCESS
# =========================================================================
add_heading_1("6. Deployment Process")

# 6.1
add_heading_2("6.1 Local Deployment (Recommended)")
add_para(
    "Local deployment is the recommended approach for FinVault AI, as it preserves the security guarantees "
    "of the air-gapped architecture. All components run directly on the user's machine."
)

add_heading_3("Prerequisites")
add_table_simple(
    ["Requirement", "Minimum", "Recommended"],
    [
        ["Python", "3.9", "3.11+"],
        ["RAM", "16 GB", "32 GB"],
        ["GPU", "Not required (CPU mode)", "8 GB+ VRAM (NVIDIA/Apple Silicon)"],
        ["Storage", "10 GB (model + app)", "20 GB (multiple models)"],
        ["LM Studio", "Latest stable", "Latest stable"],
        ["Tesseract", "5.0+", "5.5+"],
        ["OS", "macOS 12+ / Ubuntu 20.04+ / Windows 10+", "macOS 14+ / Ubuntu 22.04+"],
    ],
    col_widths=[Inches(1.4), Inches(2.2), Inches(2.8)]
)

add_heading_3("Deployment Steps")
add_para("1. Install Dependencies", bold=True)
add_code_block("pip install -r requirements.txt\nbrew install tesseract  # macOS")

add_para("2. Configure Streamlit", bold=True)
add_para("The .streamlit/config.toml file controls server settings and theme. Default configuration is provided.")

add_para("3. Start LM Studio", bold=True)
add_bullet("Launch LM Studio application")
add_bullet("Load Google Gemma 4 (or preferred model)")
add_bullet("Start the local server on port 1234")

add_para("4. Launch FinVault AI", bold=True)
add_code_block("streamlit run app.py")

add_para("5. Verify Connection", bold=True)
add_para("Navigate to the Admin Panel tab and confirm that LM Studio is connected and the model is loaded.")

add_heading_3("Deployment Process Flow")
add_code_block(
    "Install Dependencies \u2192 Configure Streamlit \u2192 Start LM Studio \u2192 Launch App \u2192 Verify Connection"
)

# 6.2
add_heading_2("6.2 Docker Deployment (Optional)")
add_para("For containerized environments, FinVault AI can be deployed using Docker:")

add_heading_3("Dockerfile Pattern")
add_code_block(
    "FROM python:3.11-slim\n"
    "\n"
    "# Install Tesseract OCR\n"
    "RUN apt-get update && apt-get install -y \\\n"
    "    tesseract-ocr \\\n"
    "    && rm -rf /var/lib/apt/lists/*\n"
    "\n"
    "WORKDIR /app\n"
    "COPY requirements.txt .\n"
    "RUN pip install --no-cache-dir -r requirements.txt\n"
    "\n"
    "COPY . .\n"
    "\n"
    "EXPOSE 8501\n"
    "CMD [\"streamlit\", \"run\", \"app.py\", \"--server.port=8501\"]"
)

add_heading_3("Docker Compose")
add_para(
    "For a complete setup, use Docker Compose to orchestrate the Streamlit application alongside "
    "a separate LM Studio container. The LM Studio container manages GPU access and model loading, "
    "while the application container handles the web UI and OCR processing."
)

# 6.3
add_heading_2("6.3 Enterprise Deployment Considerations")
add_bold_bullet("Network Isolation: ", "FinVault AI requires no egress rules. The application communicates only over localhost. Firewall rules can block all outbound traffic without affecting functionality.")
add_bold_bullet("GPU Allocation: ", "For optimal performance, dedicate GPU resources to the LM Studio process. On shared systems, use GPU isolation features (NVIDIA MPS, CUDA_VISIBLE_DEVICES) to prevent contention.")
add_bold_bullet("Multi-User Access: ", "Deploy behind a reverse proxy (Nginx, Caddy) to enable multiple users to access the same instance. Streamlit supports concurrent sessions natively.")
add_bold_bullet("Health Checks: ", "The Admin Panel provides real-time system health monitoring. For automated monitoring, implement health check endpoints that verify LM Studio connectivity and model availability.")

add_page_break()


# =========================================================================
# 7. USER DEMO GUIDE
# =========================================================================
add_heading_1("7. User Demo Guide")

# 7.1
add_heading_2("7.1 Report Generation Demo")
add_para("Follow these steps to demonstrate the AI-powered report generation workflow:")

add_para("Step 1: Load Sample Data", bold=True)
add_para(
    "In the sidebar, check the 'Use sample data' checkbox. The application loads sample_accounts.csv, "
    "which contains synthetic Q1 2026 accounting data with revenue, expense, and transfer transactions."
)

add_para("Step 2: Review Data Preview", bold=True)
add_para(
    "The data table displays all loaded records. Key metrics are computed automatically \u2014 the sample "
    "data includes approximately $401K in total credits and $349K in total debits across multiple "
    "account categories."
)

add_para("Step 3: Select Report Type", bold=True)
add_para(
    "Choose 'Executive Summary' from the report type dropdown. This generates a comprehensive overview "
    "suitable for management review."
)

add_para("Step 4: Generate Report", bold=True)
add_para(
    "Click 'Generate Report'. Gemma produces a full Q1 2026 financial report with Executive Summary, "
    "Key Metrics, Detailed Analysis, Anomalies & Risks, and Recommendations sections. The report "
    "references specific numbers from the dataset."
)

add_screenshot("01_dashboard_report.png", "Figure 4: Report Generator with sample data loaded")

# 7.2
add_heading_2("7.2 Natural Language Query Demo")
add_para("Demonstrate conversational data analysis with these steps:")

add_para("Step 1: Navigate to Data Query", bold=True)
add_para("Switch to the Data Query tab. Ensure sample data is loaded from the sidebar.")

add_para("Step 2: Ask a Question", bold=True)
add_para(
    "Type: 'What is the total revenue for Q1 2026?' The query is sent to Gemma along with the full "
    "CSV data and pre-computed statistical summary."
)

add_para("Step 3: Review the Response", bold=True)
add_para(
    "Gemma returns the answer ($401,200) with a step-by-step calculation showing how the total was "
    "derived from the data. Monetary values are formatted with proper notation."
)

add_para("Step 4: Follow-Up Questions", bold=True)
add_para(
    "Ask follow-up questions such as 'Break that down by category' or 'Which account had the highest "
    "expenses?' Chat history is maintained, allowing the model to reference previous context."
)

add_screenshot("02_data_query.png", "Figure 5: Natural Language Query with AI-powered analysis")

# 7.3
add_heading_2("7.3 OCR Document Processing Demo")
add_para("Walk through the complete OCR pipeline:")

add_para("Step 1: Navigate to Document OCR", bold=True)
add_para("Switch to the Document OCR tab and select 'Invoice' as the document type.")

add_para("Step 2: Load Sample Document", bold=True)
add_para(
    "Check 'Use sample docs' to load sample_invoice.png. Alternatively, upload your own invoice image "
    "in PNG, JPG, TIFF, BMP, or WEBP format."
)

add_para("Step 3: Extract and Parse", bold=True)
add_para(
    "Click 'Extract & Parse'. The pipeline first runs Tesseract OCR to extract raw text (typically "
    "achieving 85-95% confidence), then sends the extracted text to Gemma for structured parsing."
)

add_para("Step 4: Review Results", bold=True)
add_para(
    "The results display in two sections: (1) Raw OCR text with confidence score, and (2) Structured "
    "parsed data with labeled fields including Invoice Number, Date, Vendor, Line Items, Subtotal, "
    "Tax, and Total Amount Due."
)

add_para("Step 5: Download Results", bold=True)
add_para(
    "Download the raw extracted text or the structured parsed data for further processing. "
    "Results can be used for data entry, reconciliation, or audit documentation."
)

add_screenshot("03_ocr_processing.png", "Figure 6: OCR Pipeline \u2014 Tesseract extraction + AI parsing")

add_page_break()


# =========================================================================
# 8. ROADMAP
# =========================================================================
add_heading_1("8. Roadmap \u2014 Future Enhancements")

# 8.1
add_heading_2("8.1 Phase 2 \u2014 Intelligence (Q3 2026)")
add_bold_bullet("RAG Pipeline: ", "Integrate a local vector store (ChromaDB or FAISS) for document retrieval-augmented generation. This will enable the system to reference a library of historical documents when answering queries.")
add_bold_bullet("Multi-Document Analysis: ", "Cross-reference invoices with ledger entries, matching line items to accounting records for automated reconciliation.")
add_bold_bullet("Anomaly Detection: ", "Automated flagging of unusual transactions based on statistical outlier detection combined with LLM-based pattern analysis.")
add_bold_bullet("Auto-Categorization: ", "Use embedding similarity to automatically categorize expenses, reducing manual classification work.")

# 8.2
add_heading_2("8.2 Phase 3 \u2014 Enterprise (Q4 2026)")
add_bold_bullet("Role-Based Access Control (RBAC): ", "Define user roles (Admin, Analyst, Viewer) with permission-based access to features and data.")
add_bold_bullet("Audit Logging: ", "Comprehensive logging of all user actions, queries, and AI-generated outputs for compliance and audit trail requirements.")
add_bold_bullet("Export Formats: ", "Export reports and parsed documents to PDF and XLSX formats for distribution and archival.")
add_bold_bullet("Batch Processing: ", "Process multiple documents in a queue, enabling bulk invoice processing and batch report generation.")
add_bold_bullet("Custom Prompt Templates: ", "Allow organizations to define their own system prompts and report templates tailored to their specific financial reporting requirements.")

# 8.3
add_heading_2("8.3 Phase 4 \u2014 Scale (Q1 2027)")
add_bold_bullet("Multi-Model Support: ", "Seamlessly swap between Gemma, Llama, Mistral, and other models loaded in LM Studio without code changes.")
add_bold_bullet("Observability Dashboard: ", "Real-time monitoring of token usage, inference latency, error rates, and system resource utilization.")
add_bold_bullet("Plugin Architecture: ", "Extensible plugin system for custom data sources, report types, and processing pipelines.")
add_bold_bullet("CI/CD Pipeline: ", "Automated testing, linting, and deployment pipeline for development teams.")
add_bold_bullet("Performance Profiling: ", "Built-in profiling tools to identify bottlenecks in data processing, prompt construction, and inference workflows.")

# 8.4
add_heading_2("8.4 Security Enhancements")
add_bold_bullet("End-to-End Encryption: ", "Encrypt file uploads in transit and at rest during session processing, even though all data remains local.")
add_bold_bullet("Secure Session Management: ", "Implement token rotation and session expiration policies to prevent unauthorized access to active sessions.")
add_bold_bullet("Prompt Injection Defense: ", "Input sanitization layer to detect and prevent prompt injection attacks in user queries and uploaded documents.")
add_bold_bullet("Model Output Filtering: ", "Post-processing layer to validate and sanitize model outputs before rendering to the user interface.")

# 8.5
add_heading_2("8.5 Optimization")
add_bold_bullet("Prompt Caching: ", "Cache repeated queries and their results to reduce redundant inference calls and improve response times.")
add_bold_bullet("Streaming Responses: ", "Implement token-by-token streaming for LLM responses, providing real-time feedback during long report generation.")
add_bold_bullet("Lazy Loading: ", "Defer loading of heavy UI components until they are needed, improving initial page load performance.")
add_bold_bullet("Quantized Model Support: ", "Native support for GGUF quantized models, enabling deployment on machines with limited GPU memory.")

add_screenshot("06_roadmap.png", "Figure 7: Product Roadmap \u2014 Q2 2026 through Q1 2027")

add_page_break()


# =========================================================================
# 9. APPENDIX
# =========================================================================
add_heading_1("9. Appendix")

# 9.1
add_heading_2("9.1 System Prompts Reference")
add_para(
    "FinVault AI uses four specialized system prompts, each defined in config/settings.py. "
    "These prompts establish the AI's persona, behavioral constraints, and output format requirements "
    "for each module."
)

add_heading_3("Report Generation Prompt (SYSTEM_PROMPT_REPORT)")
add_code_block(
    "You are an expert Certified Public Accountant (CPA) and financial analyst. "
    "You generate precise, professional internal financial reports. "
    "Use proper accounting terminology (GAAP/IFRS). "
    "Structure reports with clear sections: Executive Summary, Key Metrics, "
    "Detailed Analysis, Anomalies & Risks, and Recommendations. "
    "Always reference specific numbers from the data. "
    "Format monetary values with $ signs, commas, and two decimal places."
)

add_heading_3("Data Query Prompt (SYSTEM_PROMPT_QUERY)")
add_code_block(
    "You are a financial data analyst AI. You have access to accounting data "
    "provided as CSV. Answer the user's question accurately using ONLY the data "
    "provided. Show your calculations step-by-step. "
    "Format monetary values as $X,XXX.XX. "
    "If the question cannot be answered from the data, say so clearly. "
    "If you spot anomalies or risks while answering, mention them briefly."
)

add_heading_3("OCR Document Parsing Prompt (SYSTEM_PROMPT_OCR_PARSE)")
add_code_block(
    "You are a financial document parser with expertise in invoices, work orders, "
    "purchase orders, and receipts. Extract structured data from OCR text accurately. "
    "Be precise with numbers, dates, and amounts. "
    "Return data in a clean, structured format with clear field labels. "
    "If a field is not found in the text, mark it as 'N/A'. "
    "Flag any suspicious or inconsistent values."
)

add_heading_3("OCR Vision Prompt (SYSTEM_PROMPT_OCR_VISION)")
add_code_block(
    "You are an OCR assistant. Extract ALL text visible in this document image "
    "exactly as it appears. Preserve the layout, alignment, and structure as much "
    "as possible. Include every number, date, and text element."
)

# 9.2
add_heading_2("9.2 Supported File Types")

add_heading_3("Data Files (Report Generator & Data Query)")
add_table_simple(
    ["Format", "Extension", "Library"],
    [
        ["Comma-Separated Values", ".csv", "pandas (built-in)"],
        ["Excel Workbook", ".xlsx", "pandas + openpyxl"],
        ["Legacy Excel", ".xls", "pandas + openpyxl"],
    ],
    col_widths=[Inches(2.5), Inches(1.5), Inches(2.4)]
)

add_heading_3("Image Files (Document OCR)")
add_table_simple(
    ["Format", "Extension", "Notes"],
    [
        ["PNG", ".png", "Recommended for document scans"],
        ["JPEG", ".jpg, .jpeg", "Good for photographs of documents"],
        ["TIFF", ".tiff", "Common in enterprise scanning workflows"],
        ["BMP", ".bmp", "Legacy format, full support"],
        ["WebP", ".webp", "Modern format, good compression"],
    ],
    col_widths=[Inches(1.5), Inches(1.5), Inches(3.4)]
)

# 9.3
add_heading_2("9.3 Sample Data Files")
add_para(
    "The sample_data/ directory contains seven synthetic test files for demonstrating "
    "all application features without requiring real financial data:"
)

add_table_simple(
    ["File", "Type", "Description"],
    [
        ["sample_accounts.csv", "CSV", "Q1 2026 general ledger with revenue, expense, and transfer entries"],
        ["sample_accounts_payable.csv", "CSV", "Accounts payable aging report with vendor balances"],
        ["sample_invoice.png", "Image", "Synthetic invoice image for OCR demonstration"],
        ["sample_invoice.pdf", "PDF", "Synthetic invoice document (PDF version)"],
        ["sample_purchase_order.pdf", "PDF", "Synthetic purchase order for OCR testing"],
        ["sample_receipt.png", "Image", "Synthetic receipt image for OCR demonstration"],
        ["sample_workorder.png", "Image", "Synthetic work order image for OCR testing"],
    ],
    col_widths=[Inches(2.5), Inches(0.8), Inches(3.1)]
)

add_para("")  # spacing

# Final accent line
add_accent_line()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
run = p.add_run("End of Document")
run.font.size = Pt(10)
run.font.color.rgb = MED_GRAY
run.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("FinVault AI v3.0  |  April 2026  |  Confidential")
run.font.size = Pt(9)
run.font.color.rgb = LIGHT_INDIGO


# =========================================================================
# SAVE
# =========================================================================
doc.save(OUTPUT_PATH)
print(f"Document saved to: {OUTPUT_PATH}")
